import os
import tempfile
import requests
from docx import Document

BASE = os.environ.get("BASE_URL", "http://127.0.0.1:8000/api")
YOUTUBE_URL = os.environ.get("YOUTUBE_URL")
EMAIL = "tutorctx_test@example.com"
PASSWORD = "Test1234!"
NAME = "TutorCtx Test"
ROLE = "student"


def main() -> None:
    # Register (ignore errors if already exists)
    try:
        requests.post(
            f"{BASE}/auth/register",
            json={"email": EMAIL, "password": PASSWORD, "name": NAME, "role": ROLE},
            timeout=10,
        )
    except Exception:
        pass

    login = requests.post(
        f"{BASE}/auth/login",
        json={"email": EMAIL, "password": PASSWORD},
        timeout=10,
    )
    login.raise_for_status()
    token = login.json()["token"]
    headers = {"Authorization": f"Bearer {token}"}

    # Create DOCX (use OS temp dir to avoid permission/locking issues)
    fd, path = tempfile.mkstemp(prefix="tmp_tutor_context_", suffix=".docx")
    os.close(fd)
    doc = Document()
    doc.add_paragraph(
        "Photosynthesis is the process plants use to convert light into chemical energy."
    )
    doc.add_paragraph("It typically happens in chloroplasts and produces oxygen.")
    doc.save(path)

    try:
        with open(path, "rb") as f:
            up = requests.post(
                f"{BASE}/tutor/context/upload",
                headers=headers,
                files={
                    "file": (
                        os.path.basename(path),
                        f,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                timeout=30,
            )
        up.raise_for_status()
        ctx_id = up.json()["context_id"]
    finally:
        try:
            os.remove(path)
        except OSError:
            pass

    chat = requests.post(
        f"{BASE}/tutor/chat",
        headers=headers,
        json={
            "message": "Explain the document in simple words.",
            "topic": "biology",
            "difficulty": "beginner",
            "context_ids": [ctx_id],
        },
        timeout=60,
    )
    chat.raise_for_status()

    yt_ok = None
    yt_chat = None
    if YOUTUBE_URL:
        yt_ok = requests.post(
            f"{BASE}/tutor/context/youtube",
            headers=headers,
            json={"url": YOUTUBE_URL},
            timeout=60,
        )
        yt_ok.raise_for_status()
        yt_ctx_id = yt_ok.json()["context_id"]

        yt_chat = requests.post(
            f"{BASE}/tutor/chat",
            headers=headers,
            json={
                "message": "Summarize the YouTube content in 5 bullet points.",
                "topic": "general",
                "difficulty": "beginner",
                "context_ids": [yt_ctx_id],
            },
            timeout=120,
        )
        yt_chat.raise_for_status()

    bad = requests.post(
        f"{BASE}/tutor/context/youtube",
        headers=headers,
        json={"url": "https://example.com/notyoutube"},
        timeout=30,
    )

    print("context upload ok:", up.json())
    print("chat ok (first 220 chars):", chat.json()["response"][:220])
    if yt_ok is not None:
        print("youtube ok:", yt_ok.json())
        print("youtube chat ok (first 220 chars):", yt_chat.json()["response"][:220])
    print("youtube invalid status:", bad.status_code)


if __name__ == "__main__":
    main()
